"""
Legal Translator AI - Backend API
FastAPI backend with Google Gemini and RAG pipeline using ChromaDB
"""
from  dotenv import load_dotenv
import os
import json
import hashlib
from typing import List, Dict, Optional, Any
from pathlib import Path
import shutil
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel

# Document processing
import pypdf
from docx import Document as DocxDocument
import tempfile

# LangChain and LlamaIndex
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from llama_index.core import (
    Document,
    VectorStoreIndex,
    StorageContext,
    Settings,
    ServiceContext,
)
from llama_index.core.node_parser import SimpleNodeParser
from llama_index.core.text_splitter import TokenTextSplitter
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.response_synthesizers import ResponseMode
from llama_index.embeddings.langchain import LangchainEmbedding

# ChromaDB
import chromadb
from chromadb.config import Settings as ChromaSettings

load_dotenv()

# Environment setup
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("Please set GOOGLE_API_KEY environment variable")

# Initialize FastAPI app
app = FastAPI(title="Legal Translator AI", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify exact origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize directories
UPLOAD_DIR = Path("uploads")
CHROMA_DIR = Path("chroma_db")
UPLOAD_DIR.mkdir(exist_ok=True)
CHROMA_DIR.mkdir(exist_ok=True)

# Initialize ChromaDB client
chroma_client = chromadb.PersistentClient(
    path=str(CHROMA_DIR),
    settings=ChromaSettings(anonymized_telemetry=False)
)

# Initialize Google AI models
embeddings_model = GoogleGenerativeAIEmbeddings(
    model="models/text-embedding-004",
    google_api_key=GOOGLE_API_KEY
)

llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash",
    google_api_key=GOOGLE_API_KEY,
    temperature=0.3,
    max_output_tokens=2048,
)

# Wrap embeddings for LlamaIndex
langchain_embedding = LangchainEmbedding(embeddings_model)

# Configure LlamaIndex settings
Settings.embed_model = langchain_embedding
Settings.chunk_size = 512
Settings.chunk_overlap = 50

# Storage for document indices
document_indices: Dict[str, VectorStoreIndex] = {}

# Pydantic models
class QuestionRequest(BaseModel):
    filename: str
    question: str

class ScanRequest(BaseModel):
    filename: str

class UploadResponse(BaseModel):
    success: bool
    filename: str
    message: str
    chunks_created: int

class AnswerResponse(BaseModel):
    answer: str
    sources: List[str]

class RiskClause(BaseModel):
    clause_type: str
    text: str
    explanation: str
    severity: str  # high, medium, low

class ScanResponse(BaseModel):
    risks: List[RiskClause]
    total_risks: int

# Utility functions
def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading PDF: {str(e)}")
    return text

def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading DOCX: {str(e)}")
    return text

def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading TXT: {str(e)}")

def process_document(file_path: Path, filename: str) -> str:
    """Process document based on file extension"""
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {extension}")

def create_vector_index(text: str, filename: str) -> VectorStoreIndex:
    """Create vector index from document text"""
    # Create collection name from filename (ChromaDB requirement)
    collection_name = hashlib.md5(filename.encode()).hexdigest()[:63]
    
    # Check if collection exists, delete if it does
    try:
        chroma_client.delete_collection(name=collection_name)
    except:
        pass  # Collection doesn't exist, which is fine
    
    # Create new collection
    collection = chroma_client.create_collection(
        name=collection_name,
        metadata={"filename": filename}
    )
    
    # Create ChromaDB vector store
    vector_store = ChromaVectorStore(chroma_collection=collection)
    
    # Create storage context
    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    
    # Split text into chunks
    text_splitter = TokenTextSplitter(
        chunk_size=512,
        chunk_overlap=50,
        separator="\n"
    )
    chunks = text_splitter.split_text(text)
    
    # Create documents from chunks
    documents = [Document(text=chunk) for chunk in chunks]
    
    # Create index
    index = VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        show_progress=True
    )
    
    return index

def get_or_create_index(filename: str) -> VectorStoreIndex:
    """Get existing index or raise error"""
    if filename not in document_indices:
        # Try to load from ChromaDB
        collection_name = hashlib.md5(filename.encode()).hexdigest()[:63]
        try:
            collection = chroma_client.get_collection(name=collection_name)
            vector_store = ChromaVectorStore(chroma_collection=collection)
            storage_context = StorageContext.from_defaults(vector_store=vector_store)
            index = VectorStoreIndex.from_vector_store(
                vector_store,
                storage_context=storage_context
            )
            document_indices[filename] = index
        except:
            raise HTTPException(status_code=404, detail=f"Document '{filename}' not found. Please upload it first.")
    
    return document_indices[filename]

def identify_risky_clauses(text: str) -> List[RiskClause]:
    """Use Gemini to identify risky clauses in legal document"""
    prompt = f"""
    You are a legal expert analyzing a contract or legal document. Your task is to identify potentially risky or important clauses that a non-lawyer should be aware of.
    
    Document text:
    {text[:8000]}  # Limit to avoid token limits
    
    Please identify and extract the following types of clauses if present:
    1. Indemnification/Hold Harmless clauses
    2. Limitation of Liability clauses
    3. Arbitration/Dispute Resolution clauses
    4. Termination clauses
    5. Confidentiality/Non-disclosure clauses
    6. Non-compete clauses
    7. Warranty disclaimers
    8. Payment/Fee escalation clauses
    9. Auto-renewal clauses
    10. Governing law/Jurisdiction clauses
    
    For each risky clause found, provide:
    - The type of clause
    - The exact text (keep it concise, max 200 characters)
    - A plain English explanation of what it means and why it matters
    - Severity level (high/medium/low)
    
    Return the response as a valid JSON array with objects containing: clause_type, text, explanation, severity.
    If no risky clauses are found, return an empty array [].
    
    JSON Response:
    """
    
    try:
        response = llm.invoke(prompt)
        content = response.content
        
        # Extract JSON from response
        if "```json" in content:
            json_str = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            json_str = content.split("```")[1].split("```")[0].strip()
        else:
            json_str = content.strip()
        
        # Parse JSON
        try:
            risks_data = json.loads(json_str)
            if not isinstance(risks_data, list):
                risks_data = []
        except json.JSONDecodeError:
            # Fallback to empty list if parsing fails
            risks_data = []
        
        # Convert to RiskClause objects
        risks = []
        for risk in risks_data:
            if isinstance(risk, dict):
                risks.append(RiskClause(
                    clause_type=risk.get("clause_type", "Unknown"),
                    text=risk.get("text", "")[:200],  # Limit text length
                    explanation=risk.get("explanation", "No explanation available"),
                    severity=risk.get("severity", "medium")
                ))
        
        return risks
    except Exception as e:
        print(f"Error identifying risky clauses: {str(e)}")
        return []

# API Endpoints
@app.get("/")
async def root():
    """Root endpoint"""
    return {"message": "Legal Translator AI API is running", "version": "1.0.0"}

@app.post("/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a legal document"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    # Check file extension
    allowed_extensions = ['.pdf', '.docx', '.doc', '.txt']
    file_extension = Path(file.filename).suffix.lower()
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"File type {file_extension} not supported. Allowed types: {', '.join(allowed_extensions)}"
        )
    
    # Save uploaded file
    file_path = UPLOAD_DIR / file.filename
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")
    finally:
        file.file.close()
    
    # Extract text from document
    try:
        text = process_document(file_path, file.filename)
        if not text or len(text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Document appears to be empty or unreadable")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
    
    # Create vector index
    try:
        index = create_vector_index(text, file.filename)
        document_indices[file.filename] = index
        
        # Get number of chunks created
        collection_name = hashlib.md5(file.filename.encode()).hexdigest()[:63]
        collection = chroma_client.get_collection(name=collection_name)
        num_chunks = collection.count()
        
        return UploadResponse(
            success=True,
            filename=file.filename,
            message=f"Document processed successfully. Created {num_chunks} text chunks.",
            chunks_created=num_chunks
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating vector index: {str(e)}")

@app.post("/ask", response_model=AnswerResponse)
async def ask_question(request: QuestionRequest):
    """Answer a question based on the uploaded document using RAG"""
    if not request.question or len(request.question.strip()) < 3:
        raise HTTPException(status_code=400, detail="Question is too short or empty")
    
    # Get document index
    try:
        index = get_or_create_index(request.filename)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading document: {str(e)}")
    
    # Create query engine with custom prompt
    query_engine = index.as_query_engine(
        similarity_top_k=5,
        response_mode=ResponseMode.COMPACT,
    )
    
    # Enhanced prompt for plain English explanations
    enhanced_question = f"""
    Based solely on the legal document provided, answer the following question in plain, simple English that a non-lawyer can understand.
    Avoid legal jargon where possible, and if you must use legal terms, explain them clearly.
    If the document doesn't contain information to answer the question, say so clearly.
    
    Question: {request.question}
    
    Provide a clear, concise answer:
    """
    
    try:
        # Query the index
        response = query_engine.query(enhanced_question)
        
        # Extract source nodes for transparency
        sources = []
        if hasattr(response, 'source_nodes'):
            for node in response.source_nodes[:3]:  # Limit to top 3 sources
                if hasattr(node, 'text'):
                    sources.append(node.text[:150] + "...")  # Truncate for readability
        
        return AnswerResponse(
            answer=str(response),
            sources=sources
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error querying document: {str(e)}")

@app.post("/scan", response_model=ScanResponse)
async def scan_for_risks(request: ScanRequest):
    """Scan document for risky or important clauses"""
    # Get document text
    file_path = UPLOAD_DIR / request.filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Document '{request.filename}' not found")
    
    try:
        text = process_document(file_path, request.filename)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading document: {str(e)}")
    
    # Identify risky clauses
    try:
        risks = identify_risky_clauses(text)
        return ScanResponse(
            risks=risks,
            total_risks=len(risks)
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error scanning document: {str(e)}")

@app.get("/documents")
async def list_documents():
    """List all uploaded documents"""
    documents = []
    for file_path in UPLOAD_DIR.glob("*"):
        if file_path.is_file():
            documents.append({
                "filename": file_path.name,
                "size": file_path.stat().st_size,
                "uploaded_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            })
    return {"documents": documents}

@app.delete("/documents/{filename}")
async def delete_document(filename: str):
    """Delete an uploaded document and its index"""
    file_path = UPLOAD_DIR / filename
    
    # Delete file
    if file_path.exists():
        file_path.unlink()
    
    # Delete from indices
    if filename in document_indices:
        del document_indices[filename]
    
    # Delete ChromaDB collection
    collection_name = hashlib.md5(filename.encode()).hexdigest()[:63]
    try:
        chroma_client.delete_collection(name=collection_name)
    except:
        pass  # Collection might not exist
    
    return {"message": f"Document '{filename}' deleted successfully"}

# Health check endpoint
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "api_key_set": bool(GOOGLE_API_KEY),
        "upload_dir_exists": UPLOAD_DIR.exists(),
        "chroma_dir_exists": CHROMA_DIR.exists()
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)